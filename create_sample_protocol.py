from docx import Document

def create_sample_protocol():
    doc = Document()
    doc.add_heading('Clinical Study Protocol: ARN-509-003', 0)
    
    doc.add_heading('1. Study Metadata', level=1)
    doc.add_paragraph('Protocol Number: ARN-509-003')
    doc.add_paragraph('Study Phase: Phase III')
    doc.add_paragraph('Indication: Castration-Resistant Prostate Cancer (CRPC)')
    doc.add_paragraph('Study Title: A Randomized, Double-blind, Placebo-controlled, Phase 3 Study of ARN-509 in Men with Non-metastatic Castration-resistant Prostate Cancer.')

    doc.add_heading('2. Schedule of Activities', level=1)
    doc.add_paragraph('Visit 1 (Screening): Day -28 to -1')
    doc.add_paragraph('Visit 2 (Baseline): Day 1')
    doc.add_paragraph('Visit 3: Week 4 (Day 29)')
    doc.add_paragraph('Visit 4: Week 8 (Day 57)')
    doc.add_paragraph('Visit 5: Week 12 (Day 85)')
    
    doc.add_heading('3. Assessment Procedures for Vital Signs (VS)', level=1)
    doc.add_paragraph('Vital signs will be measured at each visit (Screening, Baseline, and all subsequent treatment visits). The parameters to be collected include blood pressure (systolic and diastolic), heart rate, respiratory rate, and oral temperature.')
    doc.add_paragraph('Sitting blood pressure should be measured after the patient has been resting for at least 5 minutes.')
    
    doc.add_heading('4. Adverse Event (AE) Collection', level=1)
    doc.add_paragraph('Adverse events should be recorded throughout the study from the time of informed consent through 30 days after the last dose of study drug. Each AE should be evaluated for severity (Grade 1-5), relationship to study drug, and outcome.')
    
    doc.add_heading('5. Concomitant Medications (CM)', level=1)
    doc.add_paragraph('All concomitant medications taken by the subject during the study will be recorded. Information to be collected includes medication name, dose, frequency, route, start date, end date, and indication.')

    filename = "data/protocol/Sample_Protocol.docx"
    doc.save(filename)
    print(f"Sample protocol created at {filename}")

if __name__ == "__main__":
    create_sample_protocol()

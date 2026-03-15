from docx import Document
from docx.shared import Inches
from src.utils.state import eCRFForm, eCRFField
from typing import List

def generate_docx_report(forms: List[eCRFForm], filename: str = "eCRF_Report.docx", study_info: dict = None, visit_schedule: list = None, assessment_map: dict = None):
    doc = Document()
    doc.add_heading('Clinical eCRF Specification Report', 0)

    # 1. Study Information Section
    if study_info:
        doc.add_heading('1. Study Information', level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        for key, value in study_info.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key.replace('_', ' ').title()
            row_cells[1].text = str(value)
        doc.add_paragraph()

    # 2. Study Visit Schedule Section (Commented Out for now)
    # doc.add_heading('2. Study Visit Schedule', level=1)
    # if visit_schedule and len(visit_schedule) > 0:
    # ... (skipping logic for brevity in tool call)
    # doc.add_paragraph()

    # 3. Assessment Mapping Section (Commented Out for now)
    # if assessment_map:
    # ...
    
    # 4. Form Specification Section (Lightweight)
    doc.add_heading('2. eCRF Form Specification', level=1)

    for form in forms:
        doc.add_heading(f'Form: {form.form_name} ({form.form_id})', level=2)
        doc.add_paragraph(f"Domain: {form.domain}")
        
        if form.applicable_visits:
            doc.add_paragraph(f"Applicable Visits: {', '.join(form.applicable_visits)}")

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field ID'
        hdr_cells[1].text = 'Label'
        hdr_cells[2].text = 'Data Type'
        hdr_cells[3].text = 'Control'
        hdr_cells[4].text = 'Codelist / Options'

        for field in form.fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field.field_id
            row_cells[1].text = field.label
            row_cells[2].text = field.data_type
            row_cells[3].text = field.control_type
            
            # Format Codelist briefly
            if field.codelist:
                options = []
                # Handle possible dict structures from LLM
                for item in field.codelist:
                    if isinstance(item, dict):
                        code = item.get('code', item.get('value', ''))
                        label = item.get('label', item.get('text', ''))
                        options.append(f"{code}={label}" if code and label else str(code or label))
                    else:
                        options.append(str(item))
                row_cells[4].text = ", ".join(options)

        doc.add_page_break()

    doc.save(filename)
    return filename
